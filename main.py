import os
import shutil
import re
import logging
from collections import Counter
from tqdm import tqdm
import openpyxl
from docx import Document
import nltk
from nltk.corpus import words
import zipfile
from xml.etree import ElementTree as ET

# Hyphen normalization
HYPHENS = ["–", "—", "‐", "‒", "―"]
def normalize_hyphens(text):
    for h in HYPHENS:
        text = text.replace(h, "-")
    return text

# =====================================================
# REGEX
# =====================================================
# PART_RE — matches PART only as a STANDALONE HEADING at start of paragraph.
#
# VALID matches (standalone headings):
#   "PART I"  /  "PART - II"  /  "PART III"  /  "PART-IV"  /  "PART 1"
#   "Part I : Audit Details"   (heading with subtitle after is OK)
#
# INVALID — will NOT match (PART embedded mid-sentence):
#   "List of outstanding paras in Part-III of Inspection Report"
#   "Refer to Part II for details"
#   "As mentioned in PART I above"
#
# KEY CHANGE: ^ anchor means PART must be at the START of the paragraph text.
PART_RE = re.compile(
    r"^\s*PART[\s-]+[A-ZIVX\d]{1,4}\b",
    re.IGNORECASE | re.MULTILINE
)
OBS_RE = re.compile(
    r"\bOBS-\d+\b",
    re.IGNORECASE
)

# =====================================================
# CONFIGURATION
# =====================================================
SAVE_INTERVAL = 10


WORKING_FILE = "IR_Analysis_Working.xlsx"
VIEW_FILE = "IR_Analysis_View.xlsx"

INPUT_FOLDER = r"C:\Users\varsh\Downloads\Central offices_with only Docx\Central offices_total Docx\CNAU-AHM-01"
OUTPUT_BASE = r"nim_categorized_files"

# =====================================================
# NLTK WORD LIST
# =====================================================
try:
    english_words_set = set(words.words())
except LookupError:
    nltk.download("words")
    english_words_set = set(words.words())

# =====================================================
# LOGGING
# =====================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("ir_analyzer.log", mode="w", encoding="utf-8"),
        logging.StreamHandler()
    ]
)

# =====================================================
# SAFE DRAFT REGEX
# =====================================================
DRAFT_RE = re.compile(r'(?<![a-z])(draft|dir)(?![a-z])', re.IGNORECASE)

# =====================================================
# LANGUAGE PATTERNS
# =====================================================
LANGUAGE_PATTERNS = {
    "Hindi": re.compile(r'[\u0900-\u097F]'),
    "Malayalam": re.compile(r'[\u0D00-\u0D7F]'),
    "Tamil": re.compile(r'[\u0B80-\u0BFF]'),
    "Telugu": re.compile(r'[\u0C00-\u0C7F]'),
    "Kannada": re.compile(r'[\u0C80-\u0CFF]'),
    "Bengali": re.compile(r'[\u0980-\u09FF]'),
    "Gujarati": re.compile(r'[\u0A80-\u0AFF]'),
    "Punjabi": re.compile(r'[\u0A00-\u0A7F]')
}

ENGLISH_RE = re.compile(r'^[A-Za-z]+$')

# =====================================================
# CATEGORIES
# =====================================================
CATEGORIES = {
    "IR File Valid": "IR_File_Valid",
    "Draft IR": "Draft_IR",
    "Hindi File": "Hindi_File",
    "Malayalam File": "Malayalam_File",
    "Other Indian Language File": "Other_Indian_Language_File",
    "Other File": "Other_File",
    "Failed File": "Failed_File",
    "Same Content": "Same_Content"
}

for folder in CATEGORIES.values():
    os.makedirs(os.path.join(OUTPUT_BASE, folder), exist_ok=True)

# =====================================================
# HELPER FUNCTIONS
# =====================================================
def get_relative_path(full_path):
    return os.path.relpath(full_path, INPUT_FOLDER)

def extract_state_pr(relative_path):
    parts = relative_path.split(os.sep)
    pr = ""
    state = ""

    for i, part in enumerate(parts):
        if part.upper().startswith("PR-"):
            pr = part
            if i > 0:
                state = parts[i - 1]
            break

    return state, pr

def copy_with_pr_structure(src, category):
    """
    Copy file preserving PR folder structure inside the category folder.
    Output: OUTPUT_BASE/Category/PR-XXXXX/filename.docx
    """
    relative_path = get_relative_path(src)
    parts = relative_path.split(os.sep)

    # Find PR folder index
    pr_folder = None
    pr_index = -1
    for i, part in enumerate(parts):
        if part.upper().startswith("PR-"):
            pr_folder = part
            pr_index = i
            break

    if pr_folder and pr_index >= 0:
        # dest = OUTPUT_BASE/Category/PR-XXXXX/filename.docx
        filename = os.path.basename(src)
        dest_dir = os.path.join(OUTPUT_BASE, CATEGORIES[category], pr_folder)
        os.makedirs(dest_dir, exist_ok=True)
        dest_path = os.path.join(dest_dir, filename)
    else:
        # Fallback: no PR folder found, just put file directly in category folder
        filename = os.path.basename(src)
        dest_dir = os.path.join(OUTPUT_BASE, CATEGORIES[category])
        os.makedirs(dest_dir, exist_ok=True)
        dest_path = os.path.join(dest_dir, filename)

    shutil.copy2(src, dest_path)

def extract_text_from_corrupted_docx(file_path):
    """
    Extract text from DOCX by reading XML directly when normal extraction fails.
    """
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            xml_content = None
            try:
                xml_content = zip_ref.read('word/document.xml')
            except KeyError:
                file_list = zip_ref.namelist()
                xml_files = [f for f in file_list if 'document' in f.lower() and f.endswith('.xml')]
                if xml_files:
                    xml_content = zip_ref.read(xml_files[0])

            if not xml_content:
                return ""

            root = ET.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            text_elements = root.findall('.//w:t', ns)
            extracted_text = ' '.join([elem.text for elem in text_elements if elem.text])

            return extracted_text
    except Exception as e:
        logging.warning(f"Could not extract from XML for corrupted file: {e}")
        return ""

# =====================================================
# EXCEL WRITER
# =====================================================
class BatchExcelWriter:

    def __init__(self, filename):
        self.filename = filename
        self.workbook = openpyxl.Workbook()
        self.sheet = self.workbook.active
        self.sheet.title = "Files"
        self.metrics_sheet = self.workbook.create_sheet("Metrics")

        headers = [
            "StateCode",
            "PR_Number",
            "FileName",
            "Category",
            "EnglishPercent",
            "HindiPercent",
            "MalayalamPercent",
            "TamilPercent",
            "TeluguPercent",
            "KannadaPercent",
            "BengaliPercent",
            "GujaratiPercent",
            "PunjabiPercent",
            "PART_Found",
            "OBS_Found",
            "RelativePath"
        ]

        for col, header in enumerate(headers, 1):
            self.sheet.cell(row=1, column=col, value=header)

        self.row_num = 2

    def append_row(self, row_data):
        for col, value in enumerate(row_data, 1):
            self.sheet.cell(row=self.row_num, column=col, value=value)
        self.row_num += 1

    def update_metrics(self, metrics):
        self.metrics_sheet.delete_rows(1, 1000)
        self.metrics_sheet.cell(row=1, column=1, value="Category")
        self.metrics_sheet.cell(row=1, column=2, value="Count")

        row = 2
        for k, v in metrics.items():
            self.metrics_sheet.cell(row=row, column=1, value=k)
            self.metrics_sheet.cell(row=row, column=2, value=v)
            row += 1

    def save(self):
        self.workbook.save(self.filename)

    def close(self):
        self.workbook.close()

# =====================================================
# LANGUAGE ANALYSIS
# =====================================================
def process_text_for_language(text, lang_counter):
    for word in text.split():
        clean_word = re.sub(
            r'[^A-Za-z'
            r'\u0900-\u097F'
            r'\u0D00-\u0D7F'
            r'\u0B80-\u0BFF'
            r'\u0C00-\u0C7F'
            r'\u0C80-\u0CFF'
            r'\u0980-\u09FF'
            r'\u0A80-\u0AFF'
            r'\u0A00-\u0A7F]',
            '',
            word
        )
        if not clean_word:
            continue
        detected = False
        for lang, pattern in LANGUAGE_PATTERNS.items():
            if pattern.search(clean_word):
                lang_counter[lang] += 1
                detected = True
                break
        if not detected and ENGLISH_RE.match(clean_word):
            lang_counter["English"] += 1

def analyze_document(doc, file_path=None):
    lang_counter = Counter()
    part_found = False
    obs_found = False

    extraction_count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        extraction_count += 1
        text = normalize_hyphens(text)
        if PART_RE.search(text):
            part_found = True
        if OBS_RE.search(text):
            obs_found = True
        process_text_for_language(text, lang_counter)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    extraction_count += 1
                    text = normalize_hyphens(text)
                    if PART_RE.search(text):
                        part_found = True
                    if OBS_RE.search(text):
                        obs_found = True
                    process_text_for_language(text, lang_counter)

    try:
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                extraction_count += 1
                text = normalize_hyphens(text)
                if PART_RE.search(text):
                    part_found = True
                if OBS_RE.search(text):
                    obs_found = True
                process_text_for_language(text, lang_counter)
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                extraction_count += 1
                text = normalize_hyphens(text)
                if PART_RE.search(text):
                    part_found = True
                if OBS_RE.search(text):
                    obs_found = True
                process_text_for_language(text, lang_counter)
    except:
        pass

    if extraction_count <= 3 and file_path:
        logging.warning(
            f"EMPTY/IMAGE-ONLY FILE: '{os.path.basename(file_path)}' yielded only "
            f"{extraction_count} text item(s) via python-docx. "
            f"Attempting XML fallback recovery — this is normal for blank or image-only files."
        )
        recovered_text = extract_text_from_corrupted_docx(file_path)
        if recovered_text:
            text = normalize_hyphens(recovered_text)
            if PART_RE.search(text):
                part_found = True
            if OBS_RE.search(text):
                obs_found = True
            process_text_for_language(text, lang_counter)
            logging.info(f"✓ Successfully recovered text from {os.path.basename(file_path)}")

    total = sum(lang_counter.values())
    percentages = {}
    for lang in list(LANGUAGE_PATTERNS.keys()) + ["English"]:
        percentages[lang] = round((lang_counter[lang]*100/total), 2) if total else 0

    return percentages, part_found, obs_found

# =====================================================
# CLASSIFICATION
# =====================================================
def classify_file(filename, lang, part_found, obs_found):
    # Draft check always comes first
    if DRAFT_RE.search(filename):
        return "Draft IR"

    # --- HINDI ---
    # Any Hindi content present (even tiny amount):
    #   + has PART/OBS -> Hindi File   (it is a Hindi IR document)
    #   + no PART/OBS  -> Other File   (just stray Hindi text, not a valid IR)
    if lang["Hindi"] > 0:
        if part_found or obs_found:
            return "Hindi File"
        else:
            return "Other File"

    # --- MALAYALAM ---
    if lang["Malayalam"] > 0:
        if part_found or obs_found:
            return "Malayalam File"
        else:
            return "Other File"

    # --- OTHER INDIAN LANGUAGES ---
    other_langs = ["Tamil", "Telugu", "Kannada", "Bengali", "Gujarati", "Punjabi"]
    if any(lang[l] > 0 for l in other_langs):
        if part_found or obs_found:
            return "Other Indian Language File"
        else:
            return "Other File"

    # --- PURE ENGLISH ---
    # No non-English content at all.
    # Has PART/OBS -> IR File Valid, else Other File
    if part_found or obs_found:
        return "IR File Valid"

    return "Other File"

def is_non_english(lang):
    """Return True if the file contains any non-English Indian language content."""
    non_eng_langs = ["Hindi", "Malayalam", "Tamil", "Telugu", "Kannada", "Bengali", "Gujarati", "Punjabi"]
    return any(lang[l] > 0 for l in non_eng_langs)

# =====================================================
# WRITE ROW TO EXCEL
# =====================================================
def write_excel_row(excel_writer, state, pr, fname, category, lang, part_found, obs_found, relative_path):
    excel_writer.append_row([
        state, pr, fname, category,
        lang.get("English", 0),
        lang.get("Hindi", 0),
        lang.get("Malayalam", 0),
        lang.get("Tamil", 0),
        lang.get("Telugu", 0),
        lang.get("Kannada", 0),
        lang.get("Bengali", 0),
        lang.get("Gujarati", 0),
        lang.get("Punjabi", 0),
        "Yes" if part_found else "No",
        "Yes" if obs_found else "No",
        relative_path
    ])

# =====================================================
# ANALYZE FILE (no copy, just analyze)
# =====================================================
def analyze_file(file_path):
    """Analyze a file and return (lang, part_found, obs_found). Does NOT copy anywhere."""
    try:
        doc = Document(file_path)
        lang, part_found, obs_found = analyze_document(doc, file_path)
        return lang, part_found, obs_found, None
    except Exception as e:
        logging.error(f"FAILED analyzing | {os.path.basename(file_path)} | {e}")
        empty_lang = {l: 0 for l in list(LANGUAGE_PATTERNS.keys()) + ["English"]}
        return empty_lang, False, False, str(e)

# =====================================================
# SINGLE FILE PROCESSOR
# =====================================================
def process_single_file(file_path, excel_writer, metrics):
    """Analyze, classify, copy, and log a single file. Returns (category, lang, part_found, obs_found)."""
    fname = os.path.basename(file_path)
    relative_path = get_relative_path(file_path)
    state, pr = extract_state_pr(relative_path)

    lang, part_found, obs_found, error = analyze_file(file_path)

    if error:
        copy_with_pr_structure(file_path, "Failed File")
        metrics["Failed File"] += 1
        write_excel_row(excel_writer, state, pr, fname, "Failed File",
                        lang, part_found, obs_found, relative_path)
        return "Failed File", lang, part_found, obs_found

    category = classify_file(fname, lang, part_found, obs_found)
    copy_with_pr_structure(file_path, category)
    metrics[category] += 1
    write_excel_row(excel_writer, state, pr, fname, category,
                    lang, part_found, obs_found, relative_path)
    return category, lang, part_found, obs_found

# =====================================================
# COPY FILE AS SAME CONTENT AND LOG
# =====================================================
def send_to_samecontent(file_path, lang, part_found, obs_found, excel_writer, metrics):
    """Copy file to Same Content folder and log to excel using already-analyzed data."""
    fname = os.path.basename(file_path)
    relative_path = get_relative_path(file_path)
    state, pr = extract_state_pr(relative_path)

    copy_with_pr_structure(file_path, "Same Content")
    metrics["Same Content"] += 1
    write_excel_row(excel_writer, state, pr, fname, "Same Content",
                    lang, part_found, obs_found, relative_path)

# =====================================================
# COPY FILE AS GIVEN CATEGORY AND LOG
# =====================================================
def send_to_category(file_path, category, lang, part_found, obs_found, excel_writer, metrics):
    """Copy file to given category folder and log to excel using already-analyzed data."""
    fname = os.path.basename(file_path)
    relative_path = get_relative_path(file_path)
    state, pr = extract_state_pr(relative_path)

    copy_with_pr_structure(file_path, category)
    metrics[category] += 1
    write_excel_row(excel_writer, state, pr, fname, category,
                    lang, part_found, obs_found, relative_path)

# =====================================================
# TWO-FILE PR HANDLER
# =====================================================
def handle_two_file_pr(file_paths, excel_writer, metrics):
    """Apply case rules when a PR folder has exactly two files."""
    f1, f2 = file_paths
    fname1 = os.path.basename(f1)
    fname2 = os.path.basename(f2)

    f1_is_draft = bool(DRAFT_RE.search(fname1))
    f2_is_draft = bool(DRAFT_RE.search(fname2))

    # -------------------------------------------------------
    # CASE 4: Both files contain draft/dir
    # → copy one to Same Content, copy other to Draft IR
    # -------------------------------------------------------
    if f1_is_draft and f2_is_draft:
        logging.info(f"CASE 4 (both draft): {fname1} | {fname2}")
        lang1, part1, obs1, _ = analyze_file(f1)
        lang2, part2, obs2, _ = analyze_file(f2)
        send_to_samecontent(f1, lang1, part1, obs1, excel_writer, metrics)
        send_to_category(f2, "Draft IR", lang2, part2, obs2, excel_writer, metrics)
        return

    # Analyze both files upfront for remaining cases
    lang1, part1, obs1, err1 = analyze_file(f1)
    lang2, part2, obs2, err2 = analyze_file(f2)

    # Handle analysis failures
    if err1:
        copy_with_pr_structure(f1, "Failed File")
        metrics["Failed File"] += 1
        write_excel_row(excel_writer, *extract_state_pr(get_relative_path(f1)),
                        fname1, "Failed File", lang1, part1, obs1, get_relative_path(f1))
        lang1 = {l: 0 for l in list(LANGUAGE_PATTERNS.keys()) + ["English"]}

    if err2:
        copy_with_pr_structure(f2, "Failed File")
        metrics["Failed File"] += 1
        write_excel_row(excel_writer, *extract_state_pr(get_relative_path(f2)),
                        fname2, "Failed File", lang2, part2, obs2, get_relative_path(f2))
        lang2 = {l: 0 for l in list(LANGUAGE_PATTERNS.keys()) + ["English"]}

    if err1 and err2:
        return  # Both failed, already logged

    f1_non_english = is_non_english(lang1)
    f2_non_english = is_non_english(lang2)
    either_non_english = f1_non_english or f2_non_english

    # -------------------------------------------------------
    # CASE 3: Non-English files
    # → process both through main code normally, no Same Content
    # -------------------------------------------------------
    if either_non_english:
        logging.info(f"CASE 3 (non-English): {fname1} | {fname2}")
        if not err1:
            cat1 = classify_file(fname1, lang1, part1, obs1)
            send_to_category(f1, cat1, lang1, part1, obs1, excel_writer, metrics)
        if not err2:
            cat2 = classify_file(fname2, lang2, part2, obs2)
            send_to_category(f2, cat2, lang2, part2, obs2, excel_writer, metrics)
        return

    # -------------------------------------------------------
    # CASE 2: One file has draft/dir, other does not (both English)
    # → draft file → Same Content, other → main classification
    # -------------------------------------------------------
    if f1_is_draft and not f2_is_draft:
        logging.info(f"CASE 2 (f1 is draft): {fname1} → SameContent | {fname2} → main")
        send_to_samecontent(f1, lang1, part1, obs1, excel_writer, metrics)
        if not err2:
            cat2 = classify_file(fname2, lang2, part2, obs2)
            send_to_category(f2, cat2, lang2, part2, obs2, excel_writer, metrics)
        return

    if f2_is_draft and not f1_is_draft:
        logging.info(f"CASE 2 (f2 is draft): {fname2} → SameContent | {fname1} → main")
        send_to_samecontent(f2, lang2, part2, obs2, excel_writer, metrics)
        if not err1:
            cat1 = classify_file(fname1, lang1, part1, obs1)
            send_to_category(f1, cat1, lang1, part1, obs1, excel_writer, metrics)
        return

    # -------------------------------------------------------
    # CASE 1: No draft/dir in either filename (both English)
    # → Run both through main classification
    # → Both IR Valid           → BOTH go to IR File Valid
    # → Only one IR Valid       → IR Valid → IR File Valid, other → Same Content
    # → Neither IR Valid        → each goes to its own category
    # -------------------------------------------------------
    logging.info(f"CASE 1 (no draft, English): {fname1} | {fname2}")

    cat1 = classify_file(fname1, lang1, part1, obs1)
    cat2 = classify_file(fname2, lang2, part2, obs2)

    f1_valid = (cat1 == "IR File Valid")
    f2_valid = (cat2 == "IR File Valid")

    if f1_valid and f2_valid:
        # Both IR Valid → both go to IR File Valid
        logging.info(f"  Both IR Valid → {fname1} → IR Valid | {fname2} → IR Valid")
        send_to_category(f1, "IR File Valid", lang1, part1, obs1, excel_writer, metrics)
        send_to_category(f2, "IR File Valid", lang2, part2, obs2, excel_writer, metrics)

    elif f1_valid and not f2_valid:
        # Only f1 IR Valid → f1 → IR Valid, f2 → Same Content
        logging.info(f"  Only {fname1} is IR Valid → IR Valid | {fname2} → SameContent")
        send_to_category(f1, "IR File Valid", lang1, part1, obs1, excel_writer, metrics)
        send_to_samecontent(f2, lang2, part2, obs2, excel_writer, metrics)

    elif f2_valid and not f1_valid:
        # Only f2 IR Valid → f2 → IR Valid, f1 → Same Content
        logging.info(f"  Only {fname2} is IR Valid → IR Valid | {fname1} → SameContent")
        send_to_category(f2, "IR File Valid", lang2, part2, obs2, excel_writer, metrics)
        send_to_samecontent(f1, lang1, part1, obs1, excel_writer, metrics)

    else:
        # Neither IR Valid → each goes to its own category
        logging.info(f"  Neither IR Valid → {fname1} → {cat1} | {fname2} → {cat2}")
        send_to_category(f1, cat1, lang1, part1, obs1, excel_writer, metrics)
        send_to_category(f2, cat2, lang2, part2, obs2, excel_writer, metrics)

# =====================================================
# GROUP FILES BY PR FOLDER
# =====================================================
def group_files_by_pr(all_files):
    """
    Group files by their immediate parent folder (the PR folder).
    Returns dict: {pr_folder_path: [list of file_paths]}
    """
    pr_groups = {}
    for fp in all_files:
        pr_folder = os.path.dirname(fp)
        pr_groups.setdefault(pr_folder, []).append(fp)
    return pr_groups

# =====================================================
# MAIN
# =====================================================
if __name__ == "__main__":

    metrics = {k: 0 for k in CATEGORIES}
    excel_writer = BatchExcelWriter(WORKING_FILE)

    all_files = []
    for root, _, files in os.walk(INPUT_FOLDER):
        for f in files:
            if f.lower().endswith(".docx") and not f.startswith("~$"):
                all_files.append(os.path.join(root, f))

    logging.info(f"Found {len(all_files)} files")

    # Group files by PR folder
    pr_groups = group_files_by_pr(all_files)
    logging.info(f"Found {len(pr_groups)} PR folders")

    processed = 0

    for pr_folder, files_in_pr in tqdm(pr_groups.items(), desc="Processing PR Folders", colour="green"):

        if len(files_in_pr) == 1:
            # Single file: process normally
            process_single_file(files_in_pr[0], excel_writer, metrics)
            processed += 1

        elif len(files_in_pr) == 2:
            # Two files: apply case rules
            handle_two_file_pr(files_in_pr, excel_writer, metrics)
            processed += 2

        else:
            # More than 2 files: process each individually
            logging.warning(f"PR folder has {len(files_in_pr)} files (>2), processing each normally: {pr_folder}")
            for fp in files_in_pr:
                process_single_file(fp, excel_writer, metrics)
                processed += 1

        if processed % SAVE_INTERVAL == 0:
            excel_writer.update_metrics(metrics)
            excel_writer.save()
            try:
                shutil.copy2(WORKING_FILE, VIEW_FILE)
            except PermissionError:
                pass

    excel_writer.update_metrics(metrics)
    excel_writer.save()
    excel_writer.close()

    try:
        shutil.copy2(WORKING_FILE, VIEW_FILE)
    except PermissionError:
        pass

    logging.info("PROCESSING COMPLETE")